"""Render every final PDF page and validate the paired DOCX/PDF artifacts."""

from __future__ import annotations

import json
from pathlib import Path

import pypdfium2 as pdfium
from docx import Document
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manager_submission"
OUTPUT = ROOT / "qa" / "final_document_audit"


def render_pdf(path: Path) -> dict:
    target = OUTPUT / path.stem
    target.mkdir(parents=True, exist_ok=True)
    pdf = pdfium.PdfDocument(str(path))
    reader = PdfReader(str(path))
    page_paths = []
    page_text_lengths = []
    bad_text_pages = []
    for index in range(len(pdf)):
        page = pdf[index]
        bitmap = page.render(scale=1.7, rotation=0)
        image = bitmap.to_pil().convert("RGB")
        page_path = target / f"page_{index + 1:02d}.png"
        image.save(page_path)
        page_paths.append(page_path)
        text = reader.pages[index].extract_text() or ""
        page_text_lengths.append(len(text.strip()))
        if "\ufffd" in text or len(text.strip()) < 30:
            bad_text_pages.append(index + 1)

    thumb_width = 760
    thumbs = []
    for page_path in page_paths:
        page_image = Image.open(page_path).convert("RGB")
        ratio = thumb_width / page_image.width
        thumbs.append(page_image.resize((thumb_width, int(page_image.height * ratio))))
    rows = (len(thumbs) + 1) // 2
    cell_height = max(image.height for image in thumbs) + 54
    sheet = Image.new("RGB", (thumb_width * 2 + 72, rows * cell_height + 40), "#D9DEE5")
    draw = ImageDraw.Draw(sheet)
    font = ImageFont.truetype(r"C:\Windows\Fonts\msyh.ttc", 24)
    for index, thumb in enumerate(thumbs):
        col = index % 2
        row = index // 2
        x = 24 + col * (thumb_width + 24)
        y = 20 + row * cell_height
        sheet.paste(thumb, (x, y))
        draw.text((x + 8, y + thumb.height + 8), f"Page {index + 1}", font=font, fill="#243447")
    contact = OUTPUT / f"contact_{path.stem}.png"
    sheet.save(contact)
    return {
        "file": path.name,
        "pages": len(pdf),
        "page_text_lengths": page_text_lengths,
        "bad_text_pages": bad_text_pages,
        "all_pages_rendered": len(page_paths) == len(pdf),
        "contact_sheet": str(contact.relative_to(ROOT)),
    }


def inspect_docx(path: Path) -> dict:
    doc = Document(path)
    texts = [paragraph.text for paragraph in doc.paragraphs]
    table_text = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    joined = "\n".join([*texts, *table_text])
    return {
        "file": path.name,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "replacement_characters": joined.count("\ufffd"),
        "contains_week8": "Week 8" in joined or "WEEK 8" in joined,
        "contains_stale_release_candidate": "RELEASE CANDIDATE" in joined,
    }


def main() -> int:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    pdfs = sorted(SOURCE.glob("*.pdf"))
    docx = sorted(SOURCE.glob("*.docx"))
    report = {
        "pdf": [render_pdf(path) for path in pdfs],
        "docx": [inspect_docx(path) for path in docx],
    }
    report["passed"] = (
        len(report["pdf"]) == 4
        and len(report["docx"]) == 4
        and all(item["all_pages_rendered"] and not item["bad_text_pages"] for item in report["pdf"])
        and all(item["replacement_characters"] == 0 and item["contains_week8"] for item in report["docx"])
    )
    (OUTPUT / "document_audit.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if report["passed"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
